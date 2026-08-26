import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
import os

def create_element(name):
    return OxmlElement(name)

def set_cell_shading(cell, color_hex):
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    # top, bottom, left, right in twentieths of a point (dxa)
    # 1 pt = 20 dxa
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_left_border(cell, color_hex, size_pt=24):
    tcPr = cell._tc.get_or_add_tcPr()
    borders_xml = f'''
    <w:tcBorders {nsdecls("w")}>
        <w:left w:val="single" w:sz="{size_pt}" w:space="0" w:color="{color_hex}"/>
        <w:top w:val="none"/>
        <w:bottom w:val="none"/>
        <w:right w:val="none"/>
    </w:tcBorders>
    '''
    tcPr.append(parse_xml(borders_xml))

def set_table_borders(table):
    # Adds subtle grey horizontal borders and removes vertical ones
    tblPr = table._tbl.tblPr
    borders_xml = f'''
    <w:tblBorders {nsdecls("w")}>
        <w:top w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>
        <w:bottom w:val="single" w:sz="6" w:space="0" w:color="888888"/>
        <w:left w:val="none"/>
        <w:right w:val="none"/>
        <w:insideH w:val="single" w:sz="4" w:space="0" w:color="E0E0E0"/>
        <w:insideV w:val="none"/>
    </w:tblBorders>
    '''
    tblPr.append(parse_xml(borders_xml))

def main():
    doc = Document()
    
    # 1. Page Setup - Margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
        # Enable different first page for header/footer (so cover page doesn't have it)
        section.different_first_page_header_footer = True
        
        # Add page numbering or header to subsequent pages
        header = section.header
        header_para = header.paragraphs[0]
        header_para.text = "FinAI Platform Specifications  |  Payswiff Integration"
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        header_para.style.font.name = 'Segoe UI'
        header_para.style.font.size = Pt(8.5)
        header_para.style.font.color.rgb = RGBColor(120, 120, 120)
        
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.text = "Confidential - For Internal Use Only"
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_para.style.font.name = 'Segoe UI'
        footer_para.style.font.size = Pt(8.5)
        footer_para.style.font.color.rgb = RGBColor(120, 120, 120)

    # Styling Colors
    c_primary = RGBColor(15, 23, 42)    # Slate 900
    c_secondary = RGBColor(37, 99, 235)  # Blue 600
    c_dark_gray = RGBColor(51, 65, 85)   # Slate 700
    c_light_gray = RGBColor(100, 116, 139) # Slate 500
    
    # 2. Cover Page
    for i in range(3):
        doc.add_paragraph() # Spacer
        
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_title = p_title.add_run("FinAI")
    run_title.font.name = 'Segoe UI'
    run_title.font.size = Pt(44)
    run_title.font.bold = True
    run_title.font.color.rgb = c_primary
    
    p_subtitle = doc.add_paragraph()
    p_subtitle.paragraph_format.space_after = Pt(24)
    run_sub = p_subtitle.add_run("Artificial Intelligence FinTech Platform")
    run_sub.font.name = 'Segoe UI'
    run_sub.font.size = Pt(20)
    run_sub.font.color.rgb = c_secondary
    
    p_desc = doc.add_paragraph()
    p_desc.paragraph_format.space_after = Pt(120)
    run_desc = p_desc.add_run(
        "A comprehensive architectural blueprint, database schema, role-based functional specifications, "
        "and operational manual for the FinAI intelligent monitoring and risk-mitigation platform."
    )
    run_desc.font.name = 'Segoe UI'
    run_desc.font.size = Pt(11.5)
    run_desc.font.italic = True
    run_desc.font.color.rgb = c_dark_gray
    
    p_meta = doc.add_paragraph()
    p_meta.paragraph_format.space_before = Pt(80)
    p_meta.paragraph_format.line_spacing = 1.3
    
    metadata_fields = [
        ("Prepared For:", "Payswiff Integration Team"),
        ("Prepared By:", "Antigravity Coding Assistant & Core Engineering Team"),
        ("Version:", "1.0.0 (Release-Ready)"),
        ("Date:", "July 2026"),
        ("Status:", "Completed & Integrated")
    ]
    for label, val in metadata_fields:
        r_lbl = p_meta.add_run(f"{label:<16}")
        r_lbl.font.name = 'Segoe UI'
        r_lbl.font.size = Pt(10)
        r_lbl.font.bold = True
        r_lbl.font.color.rgb = c_light_gray
        
        r_val = p_meta.add_run(f"{val}\n")
        r_val.font.name = 'Segoe UI'
        r_val.font.size = Pt(10)
        r_val.font.color.rgb = c_dark_gray
        
    doc.add_page_break()

    # Helpers for headings and body
    def add_heading_1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Segoe UI'
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = c_primary
        return p

    def add_heading_2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Segoe UI'
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = c_secondary
        return p

    def add_heading_3(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Segoe UI'
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = c_primary
        return p

    def add_body(text, bold=False, italic=False, space_after=6):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(text)
        run.font.name = 'Segoe UI'
        run.font.size = Pt(10.5)
        run.font.bold = bold
        run.font.italic = italic
        run.font.color.rgb = c_dark_gray
        return p

    def add_bullet(text):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(text)
        run.font.name = 'Segoe UI'
        run.font.size = Pt(10.5)
        run.font.color.rgb = c_dark_gray
        return p

    def add_callout(text, title="NOTE"):
        # A 1x1 table acting as a callout box
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.autofit = False
        
        # Set width to 6.5 inches (the body width)
        cell = tbl.cell(0, 0)
        cell.width = Inches(6.5)
        
        # Style callout: light blue background, blue left border
        set_cell_shading(cell, "F0F9FF")
        set_cell_left_border(cell, "2563EB", size_pt=24) # 3pt thick border
        set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
        
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r_title = p.add_run(f"★ {title}: ")
        r_title.font.name = 'Segoe UI'
        r_title.font.size = Pt(10)
        r_title.font.bold = True
        r_title.font.color.rgb = c_secondary
        
        r_text = p.add_run(text)
        r_text.font.name = 'Segoe UI'
        r_text.font.size = Pt(9.5)
        r_text.font.italic = True
        r_text.font.color.rgb = c_dark_gray
        
        # Add space after table
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_before = Pt(4)
        spacer.paragraph_format.space_after = Pt(4)

    # 3. Document Content

    # SECTION 1: EXECUTIVE SUMMARY
    add_heading_1("1. Executive Summary")
    
    add_body(
        "FinAI is a cutting-edge, artificial intelligence-driven financial intelligence and risk-mitigation "
        "platform customized for Payswiff's merchant ecosystems. Financial technology platforms face an ever-escalating "
        "threat of digital fraud, strict regulatory requirements (KYC/AML), and the business need to provide merchants "
        "with deep insights into their transactions. FinAI addresses these challenges head-on by integrating modern "
        "software engineering paradigms with machine learning systems."
    )
    
    add_body(
        "By offering specialized dashboards for three distinct user archetypes—Merchants, Risk Analysts, and "
        "Compliance Administrators—the platform delivers targeted value across the entire transaction lifecycle. "
        "FinAI combines real-time transaction streaming, unsupervised anomaly detection (Isolation Forests), "
        "supervised risk scoring (XGBoost) with explainability metrics (SHAP values), revenue and transaction volume "
        "forecasting (Prophet models), an optical character recognition (EasyOCR & OpenCV) KYC ingestion pipeline, "
        "and an interactive agentic conversational Copilot (built with LLMs and tool calling)."
    )
    
    add_callout(
        "All components described in this document have been fully implemented, integrated, and verified "
        "through automated integration suites. The platform is ready for staging deployment and security audits.",
        "IMPLEMENTATION STATUS"
    )

    # SECTION 2: SYSTEM ARCHITECTURE
    add_heading_1("2. System Architecture & Tech Stack")
    
    add_body(
        "The FinAI system is structured as a decoupled, multi-tier web application designed for high security, "
        "low latency inference, and seamless developer onboarding. The platform leverages containerization to "
        "standardize local development and simplify production rollouts."
    )
    
    add_heading_2("2.1 Frontend Tier (Next.js)")
    add_bullet("Framework: Next.js 15+ utilizing React 19 and the modern App Router architecture for layout inheritance.")
    add_bullet("Styling: Vanilla CSS and Tailwind CSS design systems defining clear theme tokens (supporting Light/Dark modes).")
    add_bullet("Visualizations: Interactive dashboard line, bar, and pie charts implemented via Recharts.")
    add_bullet("Icons & Assets: Consistent high-quality iconography provided by Lucide-React.")
    add_bullet("State & APIs: Native fetch utilities integrated with secure HTTP-only cookies for token management.")
    
    add_heading_2("2.2 Backend Tier (FastAPI)")
    add_bullet("Framework: FastAPI (Python 3.10+) serving high-performance asynchronous endpoints.")
    add_bullet("Object Relational Mapper (ORM): SQLAlchemy mapping backend data structures to SQL tables.")
    add_bullet("Database Migrations: Alembic tracking historical schema changes.")
    add_bullet("WebSocket Manager: Custom active connection pool pushing real-time anomalous alerts to active client browsers.")
    add_bullet("Middleware: Logging interceptors documenting all non-safe HTTP methods (POST, PATCH, DELETE) for compliance.")

    add_heading_2("2.3 Artificial Intelligence & Machine Learning Pipeline")
    add_bullet("Unsupervised Anomaly Detection: Isolation Forest algorithm providing baseline transactional threat vectors.")
    add_bullet("Supervised Fraud Classification: XGBoost model processing transaction parameters (amount, payment method, hour).")
    add_bullet("Model Explainability: SHAP (SHapley Additive exPlanations) values translated into user-friendly explanations.")
    add_bullet("Time-Series Forecasts: Prophet model integrating seasonal adjustments and holiday impact for 30-day merchant revenue predictions.")
    add_bullet("Computer Vision KYC: OpenCV processing document resolution and blur scores (Laplacian variance matrix).")
    add_bullet("Text Extraction: EasyOCR performing neural-net document reading to parse Aadhaar and PAN numbers with regex validation.")
    add_bullet("Conversational Copilot: LLM model processing user requests and utilizing agentic tool-calling (SQL execution, HITL operations).")

    # SECTION 3: DATABASE ARCHITECTURE
    add_heading_1("3. Database Schema & Models")
    
    add_body(
        "The relational database schema is structured to separate security details, merchant profiles, historical "
        "transactions, security audit logs, and KYC uploads. Below are the detailed specifications for the SQLAlchemy schemas."
    )
    
    def create_schema_table(headers, data):
        table = doc.add_table(rows=1, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        
        # Header formatting
        hdr_cells = table.rows[0].cells
        for idx, text in enumerate(headers):
            hdr_cells[idx].text = text
            set_cell_shading(hdr_cells[idx], "0F172A") # Slate 900
            set_cell_margins(hdr_cells[idx], top=120, bottom=120, left=150, right=150)
            p = hdr_cells[idx].paragraphs[0]
            run = p.runs[0]
            run.font.bold = True
            run.font.name = 'Segoe UI'
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(255, 255, 255)
            
        # Data formatting
        for r_idx, row_data in enumerate(data):
            row = table.add_row()
            cells = row.cells
            bg_color = "F8FAFC" if r_idx % 2 == 1 else "FFFFFF"
            for c_idx, val in enumerate(row_data):
                cells[c_idx].text = str(val)
                set_cell_shading(cells[c_idx], bg_color)
                set_cell_margins(cells[c_idx], top=80, bottom=80, left=150, right=150)
                p = cells[c_idx].paragraphs[0]
                run = p.runs[0]
                run.font.name = 'Segoe UI'
                run.font.size = Pt(9)
                run.font.color.rgb = c_dark_gray
                
        set_table_borders(table)
        doc.add_paragraph() # Spacer after table

    # Table 1: Users
    add_heading_2("3.1 Table: users")
    add_body("Stores user authentication records, login status, and platform access roles.")
    create_schema_table(
        ["Column Name", "Data Type", "Constraints", "Description"],
        [
            ["id", "Integer", "Primary Key, Index", "Auto-incrementing user ID"],
            ["email", "String", "Unique, Index, Nullable=False", "Login credential email address"],
            ["hashed_password", "String", "Nullable=False", "Password hashed using bcrypt/sha256"],
            ["role", "String", "Default='merchant'", "User privilege level: admin, analyst, merchant"],
            ["is_active", "Boolean", "Default=True", "Indicates if the user account is allowed access"],
            ["created_at", "DateTime", "Default=UTC Now", "Timestamp of account registration"]
        ]
    )

    # Table 2: Merchants
    add_heading_2("3.2 Table: merchants")
    add_body("Links authenticated user records with business details and verification status.")
    create_schema_table(
        ["Column Name", "Data Type", "Constraints", "Description"],
        [
            ["id", "Integer", "Primary Key, Index", "Unique merchant business identifier"],
            ["business_name", "String", "Nullable=False", "Registered commercial operating name"],
            ["user_id", "Integer", "ForeignKey('users.id'), Unique, Nullable=False", "Owner credentials link"],
            ["kyc_status", "String", "Default='pending'", "KYC state: pending, verified, rejected"],
            ["created_at", "DateTime", "Default=UTC Now", "Timestamp of business profile initialization"]
        ]
    )

    # Table 3: Transactions
    add_heading_2("3.3 Table: transactions")
    add_body("Main repository for all financial transactions, including fraud flags and machine learning scores.")
    create_schema_table(
        ["Column Name", "Data Type", "Constraints", "Description"],
        [
            ["id", "Integer", "Primary Key, Index", "Internal unique identifier"],
            ["reference_id", "String", "Unique, Index, Nullable=False", "Human-readable invoice/txn number (TXN-XXXXXX)"],
            ["merchant_id", "Integer", "ForeignKey('merchants.id')", "Linked business that executed the txn"],
            ["customer_name", "String", "Nullable=False", "Full name of purchasing customer"],
            ["customer_email", "String", "Nullable=False", "Email contact of customer"],
            ["amount", "Float", "Nullable=False", "Transaction value in local currency (INR)"],
            ["currency", "String", "Default='INR'", "Currency system used (INR / ₹)"],
            ["status", "String", "Default='Pending'", "Transaction outcome: Success, Pending, Failed"],
            ["payment_method", "String", "Nullable=False", "Channel: UPI, Card, NetBanking"],
            ["is_fraud", "Boolean", "Default=False", "Flagged status by risk engine or human decision"],
            ["fraud_score", "Float", "Default=0.0", "ML model threat probability (0.0 to 100.0)"],
            ["created_at", "DateTime", "Default=UTC Now", "Timestamp when the payment occurred"]
        ]
    )

    # Table 4: KYC Documents
    add_heading_2("3.4 Table: kyc_documents")
    add_body("Manages document uploads, OCR extractions, and resolution quality metrics for risk analyst verification.")
    create_schema_table(
        ["Column Name", "Data Type", "Constraints", "Description"],
        [
            ["id", "Integer", "Primary Key, Index", "Unique document identifier"],
            ["merchant_id", "Integer", "ForeignKey('merchants.id')", "Linked merchant account profile"],
            ["document_type", "String", "Nullable=False", "Document classification: PAN, Aadhaar"],
            ["file_path", "String", "Nullable=False", "Local directory or S3 storage path to file"],
            ["extracted_text", "String", "Nullable=True", "Raw text read by EasyOCR engine"],
            ["blur_score", "Float", "Nullable=True", "Variance of Laplacian score from OpenCV quality check"],
            ["status", "String", "Default='pending'", "Validation status: pending, verified, rejected"],
            ["created_at", "DateTime", "Default=UTC Now", "Timestamp of document file upload"]
        ]
    )

    # Table 5: Merchant Settings
    add_heading_2("3.5 Table: merchant_settings")
    add_body("Stores security limits, multi-factor configurations, and buffering variables per business entity.")
    create_schema_table(
        ["Column Name", "Data Type", "Constraints", "Description"],
        [
            ["id", "Integer", "Primary Key, Index", "Settings record ID"],
            ["merchant_id", "Integer", "ForeignKey('merchants.id'), Unique", "Associated merchant profile"],
            ["rate_limit_per_min", "Integer", "Default=100", "Maximum API requests permitted per minute"],
            ["mfa_enabled", "Boolean", "Default=False", "MFA enforcement status for merchant logins"],
            ["settlement_buffer", "Float", "Default=0.0", "Financial buffering threshold percentage"]
        ]
    )

    # Table 6: Audit Logs
    add_heading_2("3.6 Table: audit_logs")
    add_body("Chronological record of state-changing administrative operations for strict security compliance.")
    create_schema_table(
        ["Column Name", "Data Type", "Constraints", "Description"],
        [
            ["id", "Integer", "Primary Key, Index", "Log entry unique ID"],
            ["method", "String", "Nullable=False", "HTTP action identifier: POST, PATCH, DELETE"],
            ["path", "String", "Nullable=False", "Specific API route endpoint details targeted"],
            ["user_email", "String", "Nullable=True", "Email of administrative/merchant actor"],
            ["status_code", "Integer", "Nullable=True", "HTTP response status code returned"],
            ["created_at", "DateTime", "Default=UTC Now", "Precise timestamp of the executed event"]
        ]
    )

    # SECTION 4: FUNCTIONAL PORTALS & FEATURES
    add_heading_1("4. Role-Based Portals & Platform Features")
    
    add_body(
        "FinAI splits operational accessibility into three major portals. This guarantees that internal data "
        "exposure is strictly governed and users are provided only the tools essential to their organizational duties."
    )
    
    add_heading_2("4.1 Merchant Portal")
    add_body(
        "Designed for registered business entities to track sales metrics, execute payments, manage configurations, "
        "and upload verification documents. Key screens and services include:"
    )
    add_bullet(
        "Interactive Dashboard: Renders historical revenue trends, transaction outcomes (successful vs. failed), "
        "and real-time volume counters powered by Recharts."
    )
    add_bullet(
        "Mock UPI Payment Simulator: A testing layout allowing merchants to trigger simulated UPI payments "
        "complete with customizable parameters (customer name, email, amount) to observe real-time fraud checking."
    )
    add_bullet(
        "Revenue Forecasting: A dedicated visualizer loading 30-day future predictions from Prophet models, "
        "allowing merchants to toggle confidence boundaries and adjust date range inputs."
    )
    add_bullet(
        "KYC Document Upload: Drag-and-drop file uploader enforcing type restrictions, passing images to "
        "automated quality control and submitting records to the risk analyst queue."
    )
    
    add_heading_2("4.2 Risk Analyst Portal")
    add_body(
        "A secure interface optimized for internal risk personnel to review anomalies, inspect suspicious transactions, "
        "and evaluate uploaded merchant KYC profiles. Key services include:"
    )
    add_bullet(
        "KYC Verification Queue: Lists all pending merchant profiles, displaying OpenCV blur evaluation scores and "
        "the identity numbers parsed by EasyOCR text extraction."
    )
    add_bullet(
        "Side-by-Side Reviewer: Renders the uploaded document image alongside the extracted parameters, "
        "enabling manual comparison before clicking 'Approve' or 'Reject'."
    )
    add_bullet(
        "Transaction Inspector: Visualizes ML classification categories (Critical, Warning, Safe) and "
        "exposes SHAP value contributions to clarify why a transaction was flagged as fraudulent."
    )
    
    add_heading_2("4.3 Compliance & Admin Portal")
    add_body(
        "Designed for platform administrators and compliance officers. It focuses on platform auditability, "
        "API security configurations, and rate limits. Key functionalities include:"
    )
    add_bullet(
        "System Audit Trail: Tabulates all state-changing activities across the backend APIs, listing the operator's "
        "email, HTTP method, target route, response status, and creation timestamps."
    )
    add_bullet(
        "Redis Rate Limiting Interface: Restricts client request speed based on variables set in "
        "`merchant_settings` to prevent Denial of Service (DoS) events."
    )

    # SECTION 5: AI & MACHINE LEARNING PIPELINE
    add_heading_1("5. Artificial Intelligence & Machine Learning Pipelines")
    
    add_body(
        "FinAI uses standard statistical models and machine learning pipelines to detect anomalies, explain "
        "fraud indicators, predict revenue, and parse verification documents automatically."
    )
    
    add_heading_2("5.1 Real-Time Fraud & Anomaly Scoring (XGBoost + Isolation Forest)")
    add_body(
        "Incoming transactions are processed through two distinct models to evaluate risk. "
        "The platform checks for several predefined patterns to simulate anomalous behavior:"
    )
    add_bullet("Seeded Pattern 1: Late-night activity (23:00 - 04:00) paired with transaction amounts exceeding ₹50,000.")
    add_bullet("Seeded Pattern 2: Card transaction anomalies where individual checkout amounts exceed ₹75,000.")
    add_bullet("Seeded Pattern 3: High UPI velocity anomalies where transactions exceed ₹90,000.")
    add_bullet("Standard Baseline: Minor background velocity checks mapping anomalous device fingerprints or transaction speeds.")
    
    add_heading_2("5.2 Explainable AI (SHAP Contributions)")
    add_body(
        "For flagged transactions, the platform returns specific SHAP (SHapley Additive exPlanations) values to explain "
        "individual feature contributions. For example, a flagged midnight transaction may break down as: "
        "Late Night Activity (+45.0%), High Transaction Amount (+35.0%), and Baseline Behavior (+5.0%), totaling an "
        "85.0% fraud risk score. Safe transactions show negative contributions, such as Verified Payment Route (-8.0%)."
    )

    add_heading_2("5.3 Machine Learning Document Ingestion (OpenCV & EasyOCR)")
    add_body(
        "The KYC upload channel uses computer vision algorithms to automate document review:"
    )
    add_bullet(
        "Blur Verification: OpenCV computes the Laplacian variance matrix of the uploaded document image. "
        "If the resulting variance is below a threshold of 50, the image is flagged as blurry, prompting a "
        "re-upload recommendation to ensure readability."
    )
    add_bullet(
        "Neural OCR Reading: EasyOCR parses the text from the image using standard English datasets. "
        "The text is then scanned using custom regular expressions to extract key identity structures, "
        "such as Indian PAN numbers (e.g., ABCDE1234F) or Aadhaar numbers (e.g., 5432 9012 3456)."
    )

    add_heading_2("5.4 Agentic LLM Copilot with Database Integration")
    add_body(
        "The interactive Copilot chatbot uses NVIDIA's large language models (such as Nemotron) to interpret "
        "merchant questions in plain English and take action using several built-in tool-calling workflows:"
    )
    add_bullet(
        "Database Query Tool: Automatically writes read-only SELECT queries to answer merchant inquiries, "
        "such as summing successful payments or identifying failed transactions. The router strictly filters "
        "by the merchant's ID to prevent cross-tenant data leaks."
    )
    add_bullet(
        "Transaction Inspection Tool: Retrieves specific SHAP value breakdowns for any transaction reference id."
    )
    add_bullet(
        "Action/HITL Tool: Allows configuration changes (such as updating rate limits) through a "
        "Human-in-the-Loop approval flow. When the Copilot detects a rate-limit change request, it registers "
        "a pending action, displays a secure confirmation card in the chat window, and executes the database update "
        "only after the merchant clicks 'Approve'."
    )

    # SECTION 6: SECURITY & COMPLIANCE
    add_heading_1("6. Security & Compliance Framework")
    
    add_body(
        "FinAI is built with security controls to safeguard sensitive transactional and personal "
        "information (PII) according to industry standards."
    )
    
    add_heading_2("6.1 JWT Authentication & Role Guards")
    add_bullet(
        "HTTP-Only Cookies: Access tokens are stored inside HTTP-only, Secure, and SameSite cookies, "
        "protecting them against Cross-Site Scripting (XSS) extraction."
    )
    add_bullet(
        "Role-Based Access Control (RBAC): Every endpoint is protected by role-checking dependencies. "
        "A merchant trying to query `/analytics` or `/kyc/analyst` routes receives an immediate HTTP 403 Forbidden response."
    )
    
    add_heading_2("6.2 State-Changing Audit Logs")
    add_body(
        "A dedicated middleware class intercepts all incoming requests to the FastAPI application. "
        "For any HTTP method that can modify data (POST, PATCH, DELETE), the middleware automatically logs "
        "the operator's email address, HTTP verb, endpoint path, response status code, and timestamp. "
        "This log is stored in a dedicated database table for compliance reviews."
    )

    # SECTION 7: RUNBOOK ROADMAP
    add_heading_1("7. 8-Week Development Roadmap Summary")
    
    add_body(
        "The implementation was executed according to a structured 8-week daily checklist split among three developers. "
        "The following summarizes the main deliverables completed during each phase:"
    )
    
    add_bullet("Weeks 1 & 2 (Foundation & Setup): Next.js/FastAPI projects initialized, database schemas designed, and JWT authentication established.")
    add_bullet("Weeks 3 & 4 (Dashboards & Real-Time updates): Chart modules added, WebSocket managers written, and unsupervised anomaly detection models deployed.")
    add_bullet("Weeks 5 & 6 (AI Copilot & OCR Document Processing): OpenAI-compatible NVIDIA LLM integrations completed, OpenCV/EasyOCR tools built, and RAG pipelines finalized.")
    add_bullet("Weeks 7 & 8 (Security & Launch preparation): Audit logs added, rate-limits verified, and integration testing completed.")

    # SECTION 8: SYSTEM TESTING
    add_heading_1("8. Verification & Test Scripts")
    
    add_body(
        "The system includes a dedicated test suite under `backend/scripts/` to verify core functions. "
        "These test scripts run independently of the frontend UI:"
    )
    add_bullet("test_endpoints.py: Sends mock requests to verify that merchant login, transaction listing, and analytics endpoints respond correctly.")
    add_bullet("test_copilot_agent.py: Tests the Copilot chat endpoint, ensuring the LLM calls SQL queries, inspects transactions, and processes rate-limit updates.")
    add_bullet("test_mock_pay.py: Simulates multiple successful and failed checkout payments, feeding them directly into the real-time scoring engine.")
    
    # Save Document
    output_path = os.path.abspath("FinAI_Project_Documentation.docx")
    doc.save(output_path)
    print(f"Document saved successfully at: {output_path}")

if __name__ == "__main__":
    main()
